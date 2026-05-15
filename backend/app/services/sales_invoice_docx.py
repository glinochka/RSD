"""Простой счёт в формате .docx для скачивания продажником."""

from __future__ import annotations

from datetime import datetime, timezone
from io import BytesIO

from docx import Document
from docx.shared import Pt

from ..alembic.models import SalesOutboundContact
from ..config import settings


def build_contact_invoice_docx(contact: SalesOutboundContact) -> bytes:
    now = datetime.now(timezone.utc).strftime("%d.%m.%Y")
    supplier = (settings.SALES_INVOICE_SUPPLIER_NAME or "Поставщик").strip()
    details = (settings.SALES_INVOICE_SUPPLIER_DETAILS or "").strip()

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)

    doc.add_heading("Счёт на оплату", level=1)
    doc.add_paragraph(f"№ {contact.id} от {now}")
    doc.add_paragraph("")
    doc.add_paragraph(f"Поставщик: {supplier}")
    if details:
        doc.add_paragraph(details)
    doc.add_paragraph("")
    doc.add_paragraph(f"Плательщик / организация: {contact.org_name or '—'}")
    if contact.lpr_name:
        doc.add_paragraph(f"Контактное лицо: {contact.lpr_name}")
    if contact.email:
        doc.add_paragraph(f"E-mail: {contact.email}")
    doc.add_paragraph("")
    doc.add_paragraph(
        "Примечание: реквизиты, сумма и основание заполните при необходимости "
        "(шаблон формируется автоматически из карточки контакта)."
    )

    table = doc.add_table(rows=2, cols=4)
    hdr = table.rows[0].cells
    hdr[0].text = "№"
    hdr[1].text = "Наименование"
    hdr[2].text = "Кол-во"
    hdr[3].text = "Сумма"
    row1 = table.rows[1].cells
    row1[0].text = "1"
    row1[1].text = f"Услуги для {contact.org_name or 'клиента'}"
    row1[2].text = "1"
    row1[3].text = "____"

    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()
